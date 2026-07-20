# -*- coding: utf-8 -*-
"""
Parse the two text docx files into a unified question JSON.
Output: scripts/_text_questions.json
"""
import json
import re
import os
import docx

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ---------- helpers ----------
OPT_RE = re.compile(r'^([A-L])[．\.\、]\s*(.+)$')  # A. / A. / A、 option line
ANS_INLINE = re.compile(r'正确答案[：:]\s*([A-L]+)')


def clean(s: str) -> str:
    return (s or '').replace('\r', '').strip()


def is_judge(options: dict) -> bool:
    """判断题: options are 正确/错误 (possibly with other text)."""
    vals = [re.sub(r'\s+', '', v) for v in options.values()]
    return any(v == '正确' or v == '错误' for v in vals)


# ---------- File A: 100道客观题 ----------
def parse_file_a():
    path = os.path.join(BASE, '100道客观题20260608.docx')
    d = docx.Document(path)
    out = []
    source_pdf = None
    section_type = None  # 'single' | 'multi'
    i = 0
    ps = d.paragraphs
    while i < len(ps):
        p = ps[i]
        style = p.style.name
        text = clean(p.text)
        # detect source PDF group heading (Heading 3 like 《...pdf》)
        if 'Heading' in style and '《' in text and '.pdf' in text:
            m = re.search(r'《(.+?)》', text)
            if m:
                source_pdf = m.group(1).replace('.pdf', '')
            i += 1
            continue
        # detect single/multi section
        if 'Heading' in style and '单选题' in text:
            section_type = 'single'
            i += 1
            continue
        if 'Heading' in style and '多选题' in text:
            section_type = 'multi'
            i += 1
            continue
        # a question paragraph: starts with "N. " and has newline options
        if text and re.match(r'^\d+\.\s', text):
            # split into lines
            lines = [ln.strip() for ln in text.split('\n') if ln.strip()]
            # first line = question stem (number + text)
            stem = re.sub(r'^\d+\.\s*', '', lines[0])
            options = {}
            for ln in lines[1:]:
                m = OPT_RE.match(ln)
                if m:
                    options[m.group(1)] = m.group(2).strip()
            # answer is the next non-empty paragraph
            ans_letters = None
            if i + 1 < len(ps):
                nxt = clean(ps[i + 1].text)
                am = ANS_INLINE.search(nxt)
                if am:
                    ans_letters = list(am.group(1))
                    i += 1  # consume answer paragraph
            if stem and options and ans_letters:
                q = {
                    'source': '100道客观题',
                    'topic': source_pdf or '',
                    'type': section_type or 'single',
                    'question': stem,
                    'options': options,
                    'answer': ans_letters,
                    'analysis': '',
                }
                out.append(q)
            i += 1
            continue
        i += 1
    return out


# ---------- File B: 全量练习题 ----------
def parse_answers(text):
    """Parse '1.B  2.C  3.ABC ...' into {1:'B',2:'C',...}"""
    res = {}
    for m in re.finditer(r'(\d+)[．\.\、]\s*([A-L]+)', text):
        res[int(m.group(1))] = list(m.group(2))
    return res


def parse_file_b():
    path = os.path.join(BASE, '企业架构与产品设计全量练习题（含标准答案）.docx')
    d = docx.Document(path)
    ps = d.paragraphs
    texts = [clean(p.text) for p in ps]

    # locate sections
    idx_single = idx_multi = idx_ans = None
    for i, t in enumerate(texts):
        if t.startswith('一、单项选择题'):
            idx_single = i
        elif t.startswith('二、多项选择题'):
            idx_multi = i
        elif t.startswith('三、全量参考答案'):
            idx_ans = i
    # answers
    single_ans, multi_ans = {}, {}
    if idx_ans is not None:
        sub_single = sub_multi = None
        for j in range(idx_ans + 1, len(texts)):
            t = texts[j]
            if t.startswith('单项选择题答案'):
                sub_single = j
            elif t.startswith('多项选择题答案'):
                sub_multi = j
        if sub_single is not None:
            blob = ' '.join(texts[sub_single + 1: sub_multi]) if sub_multi else ' '.join(texts[sub_single + 1:])
            single_ans = parse_answers(blob)
        if sub_multi is not None:
            blob = ' '.join(texts[sub_multi + 1:])
            multi_ans = parse_answers(blob)

    def parse_section(start, end, sec_type, answers):
        """Parse a questions section: 'N、stem' followed by 'A.x' option paragraphs."""
        out = []
        i = start
        while i < end:
            t = texts[i]
            m = re.match(r'^(\d+)[、\.]\s*(.+)$', t)
            if m and not OPT_RE.match(t):  # it's a question stem, not an option
                num = int(m.group(1))
                stem = m.group(2).strip()
                options = {}
                j = i + 1
                while j < end:
                    ot = texts[j]
                    om = OPT_RE.match(ot)
                    if om:
                        options[om.group(1)] = om.group(2).strip()
                        j += 1
                    elif ot == '':
                        j += 1
                        # stop if next non-empty is another question stem
                        # peek
                        k = j
                        while k < end and texts[k] == '':
                            k += 1
                        if k < end and re.match(r'^\d+[、\.]\s', texts[k]) and not OPT_RE.match(texts[k]):
                            break
                        continue
                    else:
                        break
                if stem and options:
                    ans = answers.get(num, [])
                    qtype = sec_type
                    if is_judge(options):
                        qtype = 'judge'
                    out.append({
                        'source': '全量练习题',
                        'topic': '全量汇总',
                        'type': qtype,
                        'question': stem,
                        'options': options,
                        'answer': ans,
                        'analysis': '',
                    })
                i = j
            else:
                i += 1
        return out

    single = parse_section(idx_single + 1, idx_multi, 'single', single_ans) if idx_single else []
    multi = parse_section(idx_multi + 1, idx_ans, 'multi', multi_ans) if idx_multi else []
    return single + multi


def main():
    a = parse_file_a()
    b = parse_file_b()
    # assign stable global ids
    allq = []
    for n, q in enumerate(a, 1):
        q['id'] = f'A-{n:03d}'
        allq.append(q)
    for n, q in enumerate(b, 1):
        q['id'] = f'B-{n:03d}'
        allq.append(q)

    out_path = os.path.join(BASE, 'scripts', '_text_questions.json')
    with open(out_path, 'w', encoding='utf-8') as f:
        json.dump(allq, f, ensure_ascii=False, indent=2)

    # stats
    from collections import Counter
    print(f'File A (100道客观题): {len(a)}  -> types: {Counter(q["type"] for q in a)}')
    print(f'File B (全量练习题):  {len(b)}  -> types: {Counter(q["type"] for q in b)}')
    print(f'TOTAL text questions: {len(allq)}')
    # sanity: missing answers
    missing = [q['id'] for q in allq if not q['answer']]
    if missing:
        print(f'WARNING: {len(missing)} questions missing answers: {missing[:10]}...')
    # sanity: options count
    bad_opts = [q['id'] for q in allq if len(q['options']) < 2]
    if bad_opts:
        print(f'WARNING: {len(bad_opts)} questions with <2 options: {bad_opts[:10]}...')
    print(f'Written -> {out_path}')


if __name__ == '__main__':
    main()
